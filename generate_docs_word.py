"""
AiMitra — Application Documentation Word Generator
Run: python generate_docs_word.py
Requires: pip install python-docx
Output: AiMitra_Documentation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Colours ─────────────────────────────────────────────────────────────────
PURPLE       = RGBColor(0x7C, 0x3A, 0xED)
PURPLE_LIGHT = RGBColor(0xA7, 0x8B, 0xFA)
AMBER        = RGBColor(0xD9, 0x77, 0x06)
AMBER_LIGHT  = RGBColor(0xF5, 0x9E, 0x0B)
EMERALD      = RGBColor(0x05, 0x96, 0x69)
BLUE         = RGBColor(0x25, 0x63, 0xEB)
DARK_TEXT    = RGBColor(0x1F, 0x29, 0x37)
GRAY_TEXT    = RGBColor(0x6B, 0x72, 0x80)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

HEX_PURPLE_BG  = "EDE9FE"
HEX_AMBER_BG   = "FEF3C7"
HEX_EMERALD_BG = "D1FAE5"
HEX_BLUE_BG    = "DBEAFE"
HEX_CODE_BG    = "F0FDF4"
HEX_GRAY_BG    = "F3F4F6"


# ── Helpers ──────────────────────────────────────────────────────────────────

def shade_para(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        tag = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            tag.set(qn(f"w:{k}"), v)
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_page_break(doc):
    doc.add_page_break()


class WordDoc:
    def __init__(self):
        self.doc = Document()
        # Page margins
        for section in self.doc.sections:
            section.left_margin  = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.top_margin   = Inches(0.85)
            section.bottom_margin= Inches(0.85)
        self._setup_styles()

    def _setup_styles(self):
        styles = self.doc.styles
        # Normal style
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10)
        normal.font.color.rgb = DARK_TEXT

    def h1(self, text, color=PURPLE_LIGHT):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.font.name  = "Calibri"
        run.font.size  = Pt(18)
        run.font.bold  = True
        run.font.color.rgb = color
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def h2(self, text, color=AMBER_LIGHT):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        run.font.name  = "Calibri"
        run.font.size  = Pt(13)
        run.font.bold  = True
        run.font.color.rgb = color
        return p

    def h3(self, text, color=PURPLE):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(text)
        run.font.name  = "Calibri"
        run.font.size  = Pt(11)
        run.font.bold  = True
        run.font.color.rgb = color
        return p

    def body(self, text, italic=False, color=DARK_TEXT):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name   = "Calibri"
        run.font.size   = Pt(10)
        run.font.italic = italic
        run.font.color.rgb = color
        return p

    def bullet(self, text, color=DARK_TEXT):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_after  = Pt(2)
        p.clear()
        run = p.add_run(text)
        run.font.name  = "Calibri"
        run.font.size  = Pt(10)
        run.font.color.rgb = color
        return p

    def numbered(self, text, num, color=DARK_TEXT, label_color=PURPLE_LIGHT):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(3)
        run_n = p.add_run(f"{num}.  ")
        run_n.font.bold  = True
        run_n.font.size  = Pt(10)
        run_n.font.color.rgb = label_color
        run_t = p.add_run(text)
        run_t.font.size  = Pt(10)
        run_t.font.color.rgb = color
        return p

    def code_box(self, label, code):
        # Label row
        p_label = self.doc.add_paragraph()
        p_label.paragraph_format.space_before = Pt(4)
        p_label.paragraph_format.space_after  = Pt(0)
        run = p_label.add_run(f"  {label}")
        run.font.name  = "Calibri"
        run.font.size  = Pt(8.5)
        run.font.bold  = True
        run.font.color.rgb = EMERALD
        shade_para(p_label, HEX_EMERALD_BG)

        # Code row
        p_code = self.doc.add_paragraph()
        p_code.paragraph_format.space_before = Pt(0)
        p_code.paragraph_format.space_after  = Pt(6)
        p_code.paragraph_format.left_indent  = Inches(0.15)
        run2 = p_code.add_run(code)
        run2.font.name  = "Courier New"
        run2.font.size  = Pt(8.5)
        run2.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        shade_para(p_code, HEX_CODE_BG)
        return p_code

    def kv_row(self, label, value, label_color=AMBER_LIGHT, val_color=DARK_TEXT):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.15)
        run_l = p.add_run(f"{label}  ")
        run_l.font.name  = "Calibri"
        run_l.font.size  = Pt(9)
        run_l.font.bold  = True
        run_l.font.color.rgb = label_color
        run_v = p.add_run(value)
        run_v.font.name  = "Calibri"
        run_v.font.size  = Pt(9)
        run_v.font.color.rgb = val_color
        return p

    def divider(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "A78BFA")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def tool_entry(self, icon, name, when, why, how, demo_input, demo_note,
                   header_bg=HEX_PURPLE_BG, header_color=PURPLE_LIGHT):
        # Tool header
        p_hdr = self.doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(8)
        p_hdr.paragraph_format.space_after  = Pt(1)
        run = p_hdr.add_run(f"  {icon}  {name}")
        run.font.name  = "Calibri"
        run.font.size  = Pt(11)
        run.font.bold  = True
        run.font.color.rgb = header_color
        shade_para(p_hdr, header_bg)

        self.kv_row("When:", when)
        self.kv_row("Why:", why)
        self.kv_row("How:", how)
        self.code_box("Demo Input — paste this to test:", demo_input)
        if demo_note:
            self.body(f"Expected: {demo_note}", italic=True, color=GRAY_TEXT)
        self.doc.add_paragraph()

    def stats_table(self, stats):
        table = self.doc.add_table(rows=2, cols=len(stats))
        for col_i, (val, lbl) in enumerate(stats):
            cell_v = table.cell(0, col_i)
            cell_v.paragraphs[0].clear()
            run = cell_v.paragraphs[0].add_run(val)
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = PURPLE_LIGHT
            cell_v.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shade_cell(cell_v, "F5F3FF")

            cell_l = table.cell(1, col_i)
            cell_l.paragraphs[0].clear()
            run2 = cell_l.paragraphs[0].add_run(lbl)
            run2.font.size = Pt(9)
            run2.font.color.rgb = GRAY_TEXT
            cell_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shade_cell(cell_l, "EDE9FE")
        self.doc.add_paragraph()

    def save(self, path):
        self.doc.save(path)
        print(f"Word document generated: {path}")


# ════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════
d = WordDoc()

# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
p_cover = d.doc.add_paragraph()
p_cover.paragraph_format.space_before = Pt(60)
p_cover.paragraph_format.space_after  = Pt(4)
p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_cover.add_run("AiMitra")
run_title.font.name = "Calibri"
run_title.font.size = Pt(52)
run_title.font.bold = True
run_title.font.color.rgb = PURPLE_LIGHT

p_sub = d.doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(4)
run_sub = p_sub.add_run("Your Intelligent AI Assistant")
run_sub.font.name = "Calibri"
run_sub.font.size = Pt(18)
run_sub.font.color.rgb = AMBER_LIGHT

p_tag = d.doc.add_paragraph()
p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tag.paragraph_format.space_after = Pt(24)
run_tag = p_tag.add_run("Empowering QA Engineers, Business Analysts & Developers to build better software, faster.")
run_tag.font.name   = "Calibri"
run_tag.font.size   = Pt(11)
run_tag.font.color.rgb = GRAY_TEXT
run_tag.font.italic = True

d.stats_table([("13", "AI Providers"), ("50+", "AI Models"), ("30+", "Specialised Tools"), ("4", "Workspaces")])

p_dev = d.doc.add_paragraph()
p_dev.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_dev.paragraph_format.space_before = Pt(20)
run_dev = p_dev.add_run("Developed by Shiv Kant Kumar")
run_dev.font.name  = "Calibri"
run_dev.font.size  = Pt(11)
run_dev.font.bold  = True
run_dev.font.color.rgb = PURPLE_LIGHT

p_li = d.doc.add_paragraph()
p_li.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_li = p_li.add_run("linkedin.com/in/shivkantkumar/")
run_li.font.name  = "Calibri"
run_li.font.size  = Pt(10)
run_li.font.color.rgb = BLUE

add_page_break(d.doc)

# ─────────────────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS
# ─────────────────────────────────────────────────────────────────────────────
d.h1("Table of Contents")

toc = [
    ("1.", "About AiMitra", False),
    ("2.", "Tech Stack", False),
    ("3.", "Unique Selling Points (USP)", False),
    ("4.", "Getting Started", False),
    ("5.", "Chat & Code Workspace", False),
    ("  5.1–5.7", "All Chat Modes with demo data", True),
    ("6.", "More Tools Workspace", False),
    ("  6.1–6.14", "All Dev Tools with demo data", True),
    ("7.", "JIRA Tools Workspace", False),
    ("  7.1–7.7", "All JIRA Tools with demo data", True),
    ("8.", "BA Tools Workspace", False),
    ("  8.1–8.10", "All BA Tools with demo data", True),
    ("9.", "Chrome Extension", False),
    ("10.", "AI Providers & Models", False),
    ("11.", "Architecture & Data Flow", False),
    ("12.", "About the Developer", False),
]

for num, title, is_sub in toc:
    p = d.doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.4 if is_sub else 0)
    run_n = p.add_run(f"{num}  ")
    run_n.font.name  = "Calibri"
    run_n.font.size  = Pt(10 if is_sub else 11)
    run_n.font.bold  = not is_sub
    run_n.font.color.rgb = GRAY_TEXT if is_sub else PURPLE_LIGHT
    run_t = p.add_run(title)
    run_t.font.name  = "Calibri"
    run_t.font.size  = Pt(10 if is_sub else 11)
    run_t.font.bold  = not is_sub
    run_t.font.color.rgb = GRAY_TEXT if is_sub else DARK_TEXT

add_page_break(d.doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 — ABOUT
# ─────────────────────────────────────────────────────────────────────────────
d.h1("1. About AiMitra")
d.body(
    'AiMitra ("AI Friend" in Sanskrit) is a fully-featured, open-architecture AI assistant designed for the '
    "three pillars of software delivery: Quality Assurance (QA), Business Analysis (BA), and Software "
    "Development. It consolidates more than 30 purpose-built AI tools into a single application, eliminating "
    "the need to juggle multiple AI services."
)
d.body(
    "Unlike generic AI chat tools, every workspace and every tool inside AiMitra has a carefully crafted system "
    "prompt that turns the AI into a domain expert — a senior QA lead, a seasoned BA, or a code reviewer — "
    "depending on what you need. The smart mode detector even suggests the right tool automatically when it "
    "detects keywords in your message."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 — TECH STACK
# ─────────────────────────────────────────────────────────────────────────────
d.h1("2. Tech Stack")

d.h2("Frontend")
stack_fe = [
    ("React 18.3", "UI framework with hooks, context, and automatic batching"),
    ("Tailwind CSS", "Utility-first styling with custom dark theme"),
    ("React Markdown + remark-gfm", "Renders AI responses as rich markdown with tables, code blocks"),
    ("React Syntax Highlighter", "VS Code Dark+ themed code blocks with copy button"),
    ("Axios + Fetch API", "Axios for non-streaming; native Fetch for SSE streaming"),
    ("localStorage", "Persists API keys, config, conversation history — no account needed"),
]
for name, desc in stack_fe:
    p = d.doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{name}:  ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = AMBER_LIGHT
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_TEXT

d.h2("Backend")
stack_be = [
    ("FastAPI 0.115", "High-performance async Python web framework"),
    ("Uvicorn", "ASGI server; serves both API and built React app"),
    ("httpx", "Fully async HTTP client for all AI provider API calls"),
    ("Pydantic v2", "Request/response validation with strict typing"),
    ("SSE-Starlette", "Server-Sent Events for real-time token streaming"),
    ("BeautifulSoup4 + lxml", "HTML parsing for the web page DOM analyser"),
    ("Supabase", "Optional auth/database integration"),
]
for name, desc in stack_be:
    p = d.doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{name}:  ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = PURPLE_LIGHT
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_TEXT

d.h2("AI Integration")
d.body(
    "AiMitra supports 13 AI providers (OpenAI, Anthropic, Google Gemini, Groq, Mistral, DeepSeek, xAI/Grok, "
    "Together AI, Perplexity, Cerebras, OpenRouter, Fireworks AI, Cohere) and 50+ individual models. "
    "API keys are entered by the user, stored in localStorage, and forwarded per-request. They are never "
    "stored on the server. All providers support both streaming (SSE) and non-streaming modes."
)

add_page_break(d.doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3 — USP
# ─────────────────────────────────────────────────────────────────────────────
d.h1("3. Unique Selling Points (USP)")

usps = [
    ("All-in-One Platform",
     "30+ specialised tools across 4 workspaces (Chat, More Tools, JIRA, BA). No context switching between "
     "multiple AI services. Everything a QA, BA, and Dev needs in one interface."),
    ("No Vendor Lock-in — 13 Providers",
     "Switch between OpenAI, Gemini, Groq (free), Anthropic, DeepSeek and 9 more at any time. Free-tier "
     "providers (Groq, Gemini, Cerebras, OpenRouter) mean you can use the full feature set at zero cost."),
    ("Domain-Expert AI (Specialised Prompts)",
     "Each tool has a meticulously crafted system prompt that transforms the AI into a domain expert: senior "
     "QA lead for test plans, seasoned BA for BRDs, expert debugger for code issues."),
    ("Smart Mode Detector",
     "Paste a stack trace in the wrong mode? AiMitra detects it and suggests Code Debugging. Ask for test cases "
     "in chat? It suggests Test Case Generator. Uses keyword analysis with no extra API calls."),
    ("Real-time Streaming",
     "Every AI response streams token-by-token to the UI. No waiting for the full response."),
    ("Chrome Extension",
     "The built-in Chrome extension reads the live DOM of any webpage, captures real CSS selectors and XPaths, "
     "and generates test code grounded in the actual page — not generic placeholders."),
    ("Conversation History",
     "Every chat and tool output is auto-saved to local history with timestamp, model, and title. Resume any "
     "previous conversation with a single click. Stored in your browser only."),
    ("Zero Infrastructure Cost",
     "No database required. No user accounts. No monthly subscription. Deploy on any free hosting "
     "(Render, Railway, Vercel) or run locally with two terminal commands."),
    ("BA Tools — Often Overlooked",
     "Most AI tools focus on developers. AiMitra's dedicated BA workspace with User Story Generator, "
     "Gap Analysis, BRD Generator, Impact Analysis makes it the only QA/dev AI tool that also fully serves BAs."),
    ("Comparison with Alternatives",
     "vs. ChatGPT/Claude.ai: no specialised QA/BA prompts, no JIRA integration, no Chrome extension. "
     "vs. GitHub Copilot: code only, no QA, no BA, no test planning. "
     "vs. Testim/Mabl: expensive, no BA tools, no multi-provider AI. "
     "AiMitra: free-tier capable, multi-role, multi-provider, all-in-one."),
]

for title, desc in usps:
    d.h3(f"* {title}", AMBER_LIGHT)
    d.body(desc)

add_page_break(d.doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 — GETTING STARTED
# ─────────────────────────────────────────────────────────────────────────────
d.h1("4. Getting Started")

d.h2("Prerequisites")
for item in ["Node.js 18+", "Python 3.11+", "Git",
             "A free API key from Groq (console.groq.com) or Google AI Studio"]:
    d.bullet(item)

d.h2("Installation — Local Development")
d.code_box("Step 1 — Clone the repository",
           "git clone https://github.com/shivkantkumar/aimitra.git\ncd aimitra")
d.code_box("Step 2 — Backend setup",
           "cd backend\npip install -r requirements.txt\ncp .env.example .env\nuvicorn main:app --reload --port 8000")
d.code_box("Step 3 — Frontend (new terminal)",
           "cd frontend\nnpm install\nnpm start    # opens http://localhost:3000")
d.code_box("Step 4 — Add your API key",
           "Open the app in browser -> Sidebar -> Provider: Groq -> API Key: paste key")

d.h2("First Use Walkthrough")
steps_start = [
    "Select a Provider in the left sidebar (Groq is free — no credit card required).",
    "Paste your API key and press Enter. A green checkmark confirms it is saved.",
    "The default workspace is Chat. Type any message and press Enter or click Send.",
    'Try: "Write test cases for a login page" — the smart detector suggests Test Case Generator.',
    "Switch to More Tools (top nav) to access Code Review, BDD Generator, API Test Generator, and more.",
    "Switch to BA Tools to access User Story Generator, BRD Generator, Gap Analysis, and more.",
    "Switch to JIRA to create tickets, validate stories, or generate JQL queries.",
    "Install the Chrome Extension from the chrome-extension/ folder to generate tests from any webpage.",
]
for i, s in enumerate(steps_start, 1):
    d.numbered(s, i)

d.h2("Configuration Options")
configs = [
    ("Provider", "Choose from 13 AI providers. Switch at any time."),
    ("Model", "Select any model from the chosen provider's list."),
    ("Temperature", "0.0 = precise/deterministic. 1.0 = creative/varied. Default: 0.7."),
    ("Streaming", "Toggle real-time token streaming on/off."),
    ("API Key", "Stored in browser localStorage. Never sent to any server except the chosen AI provider."),
]
for name, desc in configs:
    p = d.doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{name}:  ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = AMBER_LIGHT
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_TEXT

add_page_break(d.doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5 — CHAT WORKSPACE
# ─────────────────────────────────────────────────────────────────────────────
d.h1("5. Chat & Code Workspace")
d.body(
    "The Chat workspace provides 7 specialised modes, each with a tailored AI system prompt. "
    "The Smart Mode Detector automatically suggests the best mode based on your message. "
    "All conversations are auto-saved to Activity History."
)

chat_tools = [
    ("💬", "Text / Code Generation",
     "Everyday use — answering questions, generating code, writing documentation.",
     "Senior engineer persona. Provides structured, well-explained answers.",
     "Select 'Text/Code Generation' from sidebar -> type your question -> press Enter.",
     "Write a Python function to validate email addresses using regex. Add type hints and docstring.",
     "Returns a clean Python function with type hints, docstring, regex pattern, and usage example."),

    ("🐛", "Code Debugging",
     "When you have an error, stack trace, or broken code that needs fixing.",
     "Expert debugger — identifies root cause, explains the bug, provides a corrected fix.",
     "Paste your code + the error message -> select Code Debugging mode -> submit.",
     "TypeError: unsupported operand type(s) for +: 'int' and 'str'\nLine 15: result = count + \" items\"\nFix this.",
     "Identifies type mismatch, explains root cause, provides fixed code with str(count)."),

    ("🧪", "Test Case Generator",
     "When you need Selenium/Playwright test cases for a feature or user story.",
     "Generates production-quality tests with assertions, waits, POM structure, and edge cases.",
     "Describe the feature or paste the user story -> select Test Case Generator -> submit.",
     "Generate Playwright Python test cases for a login page with username, password, and submit button. Include happy path, wrong password, empty fields, and SQL injection tests.",
     "Returns complete Playwright test file with fixtures and 4 test functions."),

    ("📋", "Test Plan Generator",
     "At the start of a sprint when you need a formal QA test plan.",
     "Creates IEEE 829-compliant test plans with scope, risks, entry/exit criteria.",
     "Select Test Plan Generator -> describe the feature/project -> submit.",
     "Create a test plan for an e-commerce checkout flow: cart, address, payment (Stripe), and order confirmation. API + UI testing required.",
     "Returns structured test plan with objectives, scope, test types, risks, schedule."),

    ("🎯", "DOM Locator Generator",
     "When writing automation scripts and you need reliable XPath/CSS selectors.",
     "Generates robust, stable locators using ID, data-testid, aria-label — not fragile nth-child paths.",
     "Describe the element or paste the HTML snippet -> select DOM Locator Generator -> submit.",
     '<button class="btn-primary" id="submit-order" aria-label="Place Order">Place Order</button>',
     "Returns CSS selector, XPath, and Playwright/Selenium code snippets."),

    ("🔍", "Web Search",
     "When you need current information — latest library versions, recent CVEs, today's news.",
     "Uses the AI's broad knowledge to answer time-sensitive questions with references.",
     "Select Web Search mode -> type your query -> submit.",
     "What are the latest breaking changes in Playwright 1.44 and how do they affect existing tests?",
     "Returns a summary of breaking changes, migration guide, and links to the official changelog."),

    ("🎨", "Image Generation",
     "When you need to describe or prompt an image for design, wireframing, or documentation.",
     "Returns detailed image prompts for DALL-E/Midjourney since direct generation depends on provider.",
     "Select Image Generation mode -> describe what you want -> submit.",
     "Create a logo for a software testing tool called 'AiMitra'. Modern, dark theme, violet and blue gradient, robot icon.",
     "Returns a detailed DALL-E/Stable Diffusion prompt plus design specifications."),
]

for icon, name, when, why, how, demo, note in chat_tools:
    d.tool_entry(icon, name, when, why, how, demo, note, HEX_PURPLE_BG, PURPLE_LIGHT)

add_page_break(d.doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6 — MORE TOOLS
# ─────────────────────────────────────────────────────────────────────────────
d.h1("6. More Tools Workspace", EMERALD)
d.body(
    "The More Tools workspace contains 14 specialised one-shot tools for developers. Each tool has a dedicated UI "
    "with options (language, framework, pattern), a focused system prompt, and a copy/download button for the output."
)

dev_tools = [
    ("🧩", "Chrome Extension",
     "When you need to generate test code for a live webpage directly from your browser.",
     "Reads the actual DOM — real CSS selectors, XPaths, form structures — and feeds them to AI.",
     "Install from chrome-extension/ folder -> open any webpage -> click extension icon -> Generate.",
     "Open https://automationpractice.pl/ in Chrome, click the extension, pick Playwright + Python + POM, click Generate.",
     "Returns a complete POM class + test file using real selectors from the live page."),

    ("🧭", "Tool Helper",
     "When you are not sure which AiMitra tool to use for your task.",
     "Guides you to the right tool by matching your goal to the best workspace/tool.",
     "Select Tool Helper -> describe what you want to achieve -> get directed to the right tool.",
     "I want to test a REST API that returns user profiles. What tool should I use and how?",
     "Recommends API Test Generator, explains how to use it, shows a sample cURL input format."),

    ("🔍", "Code Explainer",
     "When you receive unfamiliar code and need to understand it quickly.",
     "Breaks down complex code into plain English with line-by-line or block explanations.",
     "Paste the code snippet -> click Explain -> choose detail level.",
     "def memoize(func):\n    cache = {}\n    def wrapper(*args):\n        if args not in cache:\n            cache[args] = func(*args)\n        return cache[args]\n    return wrapper",
     "Explains the memoization decorator pattern, closure, and cache dictionary with a real usage example."),

    ("🕵️", "Code Review",
     "Before submitting a pull request or reviewing a colleague's code.",
     "Reviews for bugs, security issues, performance, readability, and SOLID principles.",
     "Paste the code -> select language -> click Review.",
     "def get_user(id):\n    query = f\"SELECT * FROM users WHERE id={id}\"\n    return db.execute(query)",
     "Flags SQL injection vulnerability, suggests parameterized queries, gives security score 2/10."),

    ("🐛", "Debug & Fix",
     "When you have broken code and want the AI to identify and fix the bug automatically.",
     "Combines root cause analysis with a working fix. Explains what changed and why.",
     "Paste the broken code + the error -> click Fix.",
     "async function fetchData() {\n  const data = await fetch('/api/users')\n  return data.json  // missing parentheses\n}",
     "Identifies missing () on .json, returns corrected code with explanation."),

    ("🔄", "Code Converter",
     "When migrating code between languages, frameworks, or syntax versions.",
     "Accurately converts logic while adapting idioms — e.g. Python list comprehensions to Java streams.",
     "Paste source code -> select source language -> select target language -> convert.",
     "Convert this Python to JavaScript:\nusers = [u for u in all_users if u['active'] and u['age'] >= 18]",
     "Returns: const users = allUsers.filter(u => u.active && u.age >= 18); with explanation."),

    ("📝", "Regex Builder",
     "When you need to write or understand a regular expression for validation or parsing.",
     "Generates regex with explanation, test cases, and code snippet for your chosen language.",
     "Describe the pattern in plain English -> select language -> build.",
     "Regex to validate an Indian mobile number: starts with 6, 7, 8, or 9, exactly 10 digits total.",
     "Returns: /^[6-9]\\d{9}$/ with explanation, Python/JS usage code, and 5 test cases."),

    ("🗄️", "SQL Helper",
     "When writing complex SQL queries, optimising slow queries, or designing schema.",
     "Writes optimised SQL with proper JOINs, indexes, window functions, and CTEs.",
     "Describe what data you need in plain English -> select SQL dialect -> generate.",
     "Get the top 5 customers by total order value in the last 90 days, include their email and order count. MySQL.",
     "Returns a CTE-based query with JOIN, GROUP BY, ORDER BY, LIMIT, and clause explanation."),

    ("📦", "Git Assistant",
     "When you need help with git commands, commit messages, branching, or resolving conflicts.",
     "Provides exact git commands, explains the effect, and warns about destructive operations.",
     "Describe your git situation -> get commands and explanation.",
     "I accidentally committed my .env file with API keys to main. How do I remove it from git history completely?",
     "Returns: git filter-branch or BFG Repo Cleaner steps + force push + .gitignore fix."),

    ("🐳", "DevOps Generator",
     "When setting up CI/CD pipelines, Dockerfiles, GitHub Actions, or Kubernetes manifests.",
     "Generates production-ready DevOps configs tailored to your stack and environment.",
     "Select the DevOps artefact type -> describe your stack -> generate.",
     "Generate a GitHub Actions CI/CD pipeline for a React + FastAPI app: lint, test, build Docker image, push to DockerHub, deploy to Render.",
     "Returns a complete .github/workflows/deploy.yml with jobs, secrets usage, and Docker build/push steps."),

    ("🎲", "JSON & Mock Data",
     "When you need sample data for testing, prototyping, or API mocking.",
     "Generates realistic, schema-consistent JSON mock data with proper data types and relationships.",
     "Describe the data structure or paste a schema -> specify count -> generate.",
     "Generate 5 mock user records with: id (UUID), name, email, role (admin/user/viewer), createdAt (ISO date), isActive.",
     "Returns a JSON array of 5 realistic records with proper UUIDs, plausible names, and varied roles."),

    ("🥒", "BDD Generator",
     "When writing Gherkin feature files for Cucumber, SpecFlow, Behave, or Pytest-BDD.",
     "Creates complete .feature files with Background, Scenario Outlines, Examples tables, and step stubs.",
     "Paste the user story -> select framework -> select scenario coverage -> generate.",
     "As a registered user, I want to reset my password via email link, so that I can regain account access. Include edge cases: expired link, already used link, invalid email.",
     "Returns a .feature file with 5 scenarios + Scenario Outline + step definition stubs."),

    ("🔌", "API Test Generator",
     "When writing automated tests for REST APIs using Postman, pytest, or RestAssured.",
     "Generates complete API test suites with setup, authentication, status checks, schema validation, edge cases.",
     "Paste the API endpoint details -> select framework -> generate.",
     "POST /api/auth/login -- body: {email, password} -- returns JWT token. Generate pytest tests for success, wrong password, missing fields, rate limiting.",
     "Returns a complete pytest file with requests library, fixtures, parametrize, and 4 test functions."),

    ("♿", "A11y Checker",
     "When reviewing a UI for accessibility compliance (WCAG 2.1 AA) before release.",
     "Analyses HTML/component code for missing ARIA labels, contrast issues, keyboard navigation.",
     "Paste your HTML or React component -> click Check -> get prioritised issues with fixes.",
     "<button onclick='submit()' style='color:#aaa; background:#bbb'>Submit</button>",
     "Flags: missing type attribute, insufficient contrast (3.1:1 vs required 4.5:1), no aria-label, no focus style."),
]

for icon, name, when, why, how, demo, note in dev_tools:
    d.tool_entry(icon, name, when, why, how, demo, note, HEX_EMERALD_BG, EMERALD)

add_page_break(d.doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 7 — JIRA TOOLS
# ─────────────────────────────────────────────────────────────────────────────
d.h1("7. JIRA Tools Workspace", RGBColor(0x93, 0xC5, 0xFD))
d.body(
    "The JIRA workspace contains 7 AI-powered tools designed to make Jira work faster and more consistent. "
    "Every tool follows Agile best practices and outputs content ready to copy-paste into Jira."
)

jira_tools = [
    ("🤖", "Ask Rovo",
     "When you have a general Jira/Agile question and want expert guidance.",
     "Acts as a Jira expert — explains workflows, best practices, and Agile ceremonies.",
     "Switch to JIRA workspace -> Ask Rovo -> type your question.",
     "What is the difference between a Story, Task, and Sub-task in Jira? When should I use each?",
     "Returns a clear comparison table with use cases and examples for each issue type."),

    ("🎫", "Ticket Creator",
     "When you need a well-structured Jira Story or Epic from a rough idea.",
     "Produces a complete ticket with Summary, Description, Acceptance Criteria, Story Points, Labels.",
     "Describe the feature in a few sentences -> click Create Ticket -> copy the output into Jira.",
     "We need to add a 'Remember Me' checkbox to the login page. It should keep the user logged in for 30 days using a secure HTTP-only cookie.",
     "Returns a complete Jira ticket with summary, description, 3 acceptance criteria, and 3 story points."),

    ("🐛", "Bug Creator",
     "When logging a bug found during testing and you need a thorough, reproducible report.",
     "Creates a professional bug report with Environment, Steps to Reproduce, Expected vs Actual, Severity.",
     "Describe the bug -> click Create Bug -> copy to Jira.",
     "On Safari iOS 17, the checkout button becomes unresponsive after selecting a delivery address. Affects 100% of iOS Safari users.",
     "Returns a bug ticket with severity Critical, 4 reproduction steps, and suggested fix area."),

    ("🔍", "JQL Search",
     "When you need to find specific issues in Jira using complex criteria.",
     "Generates optimised JQL queries with explanations and alternative queries for edge cases.",
     "Describe what issues you want to find -> click Generate JQL -> paste into Jira board filter.",
     "Find all critical bugs assigned to me that were created in the last 2 weeks and are still open, sorted by creation date.",
     "assignee = currentUser() AND priority = Critical AND issuetype = Bug AND created >= -14d AND status != Done ORDER BY created DESC"),

    ("📋", "Test Plan Review",
     "When you have a test plan in Jira and want an AI review for completeness and quality.",
     "Reviews test plans for missing coverage areas, unclear acceptance criteria, risk gaps.",
     "Paste the test plan content -> click Review -> get a structured feedback report.",
     "Test Plan: Login feature. Scope: UI testing only. Tests: valid login, invalid login. Exit criteria: all tests pass.",
     "Flags: missing API test scope, no performance/security testing, vague exit criteria."),

    ("✅", "Ticket Validator",
     "Before a ticket enters a sprint — to check if it is ready (Definition of Ready).",
     "Validates against Agile best practices: clear acceptance criteria, estimated, no ambiguous language.",
     "Paste the ticket content -> click Validate -> get a readiness score and list of issues.",
     "Story: As a user I want a better dashboard. AC: Dashboard should be improved. SP: TBD. Priority: High.",
     "Score 3/10. Issues: too vague, no story points, acceptance criteria not testable."),

    ("💬", "Comment Generator",
     "When you need to add a professional update comment to a Jira ticket.",
     "Generates clear, professional comments for status updates, blockers, QA sign-offs, deployments.",
     "Describe the update -> select comment type -> generate.",
     "QA has completed testing of the login feature. Found 1 minor bug (misaligned button on mobile -- logged as PROJ-456). All critical and major tests passed.",
     "Returns a formatted QA sign-off comment with test summary, linked bug, and recommendation."),
]

for icon, name, when, why, how, demo, note in jira_tools:
    d.tool_entry(icon, name, when, why, how, demo, note, HEX_BLUE_BG, RGBColor(0x93, 0xC5, 0xFD))

add_page_break(d.doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 8 — BA TOOLS
# ─────────────────────────────────────────────────────────────────────────────
d.h1("8. BA Tools Workspace", AMBER_LIGHT)
d.body(
    "The BA (Business Analyst) workspace is AiMitra's most unique differentiator — no other QA/dev AI tool "
    "offers a dedicated BA toolkit. 10 tools covering the full BA lifecycle from requirements gathering to "
    "stakeholder communication."
)

ba_tools = [
    ("📖", "User Story Generator",
     "At the start of a sprint when converting stakeholder requirements into development-ready user stories.",
     "Produces stories in standard 'As a... I want... So that...' format with AC, story points, and DoD.",
     "Select User Story Generator -> choose the user role -> select story type -> describe the feature.",
     "Feature: Users should be able to export their dashboard data as a CSV file. Role: Data Analyst.",
     "Returns a complete user story with title, narrative, 4 acceptance criteria, 5 story points, and DoD checklist."),

    ("✅", "Acceptance Criteria",
     "After writing a user story when you need detailed, testable acceptance criteria.",
     "Generates criteria in Gherkin (Given/When/Then), checklist, or both formats — ready for QA handoff.",
     "Paste the user story -> select format (Gherkin/checklist/both) -> generate.",
     "As a user, I want to receive an email notification when my order ships, so I can track my delivery.",
     "Returns 5 Gherkin scenarios covering: order ships, email content, multiple items, cancelled order, unsubscribe."),

    ("🎭", "Use Case Generator",
     "During requirements analysis when you need formal use case documentation for complex features.",
     "Creates complete use cases with Actors, Preconditions, Main Success Scenario, Alternative and Exception Flows.",
     "Describe the feature -> list the actors -> select complexity -> generate.",
     "Feature: Online appointment booking system. Actors: Patient, Doctor, Receptionist, Notification Service. Complexity: detailed.",
     "Returns UC-001 with 8-step main flow, 3 alternative flows, and exception handling."),

    ("🔎", "Requirements Analyzer",
     "When reviewing a requirements document before development — to catch issues early.",
     "Identifies gaps, ambiguities, conflicts, and testability issues. Each finding cites original text.",
     "Paste the requirements text -> select analysis types (gaps/ambiguities/conflicts/testability) -> analyze.",
     "REQ-01: The system shall be fast. REQ-02: Users can upload files. REQ-03: The dashboard should show relevant data.",
     "Flags: REQ-01 has no measurable metric (suggest: <2s), REQ-02 no size/format constraint, REQ-03 ambiguous."),

    ("🔄", "Process Flow Generator",
     "When documenting a business process for sign-off, training material, or system design.",
     "Creates numbered process flows with decision points, swim lanes by role, and Start/End markers.",
     "Describe the process -> select AS-IS/TO-BE/Both -> toggle swim lanes -> generate.",
     "Document the employee expense reimbursement process: employee submits receipt, manager approves/rejects, finance processes payment, employee receives funds.",
     "Returns a 4-lane swim-lane flow (Employee/Manager/Finance/System) with 12 numbered steps and 2 decision points."),

    ("📄", "BRD Generator",
     "When preparing formal project documentation for executive sign-off or vendor procurement.",
     "Generates professional BRD sections with numbered requirements (REQ-001 format), MoSCoW priority.",
     "Describe the project/feature -> select section (full/scope/functional/non-functional/constraints) -> generate.",
     "Project: Customer Self-Service Portal. Customers can view invoices, raise support tickets, and track order status. Section: Full BRD.",
     "Returns a complete BRD with 15+ numbered requirements, priority levels, constraints, and assumptions."),

    ("📊", "Gap Analysis",
     "When planning a system migration, process improvement, or digital transformation initiative.",
     "Creates a structured gap table (Gap | Category | Impact | Priority | Recommendation) and phased roadmap.",
     "Describe current state in first box -> desired state in second box -> select category -> analyze.",
     "Current: Manual invoice processing via email, 3-day turnaround, 15% error rate.\nDesired: Automated invoice portal, same-day processing, <1% error rate.",
     "Returns 6 identified gaps across Process/Technology/People categories and a 3-phase roadmap."),

    ("📧", "Stakeholder Update",
     "After a milestone, incident, or scope change when stakeholders need to be informed professionally.",
     "Generates tailored communications — executive brevity for C-suite, technical detail for dev teams.",
     "Describe the update -> select audience -> select tone -> generate.",
     "The mobile app launch has been delayed by 2 weeks due to a critical payment gateway integration issue discovered in UAT. New launch date: March 15.",
     "Returns an executive email with subject line, 3-bullet summary, business impact, and next steps."),

    ("📝", "Meeting Summarizer",
     "After any meeting — requirements workshop, sprint review, stakeholder call — to capture outcomes.",
     "Extracts Key Decisions, Action Items (with owner/due date), Discussion Points, and Parking Lot items.",
     "Paste raw meeting notes -> select meeting type -> summarize.",
     "Discussed login redesign, John said we need SSO, Sarah mentioned deadline is Q2, Mike will check with security team by Friday, budget not confirmed yet.",
     "Returns structured summary: 1 decision, 2 action items (Mike: security review by Friday), 1 parking lot."),

    ("💥", "Impact Analysis",
     "Before approving a change request or architectural decision that affects multiple teams or systems.",
     "Assesses impact across People, Process, Technology, and Data — with impact level and mitigation strategies.",
     "Describe the proposed change -> select impact areas -> analyze.",
     "We are replacing our MySQL database with MongoDB for the product catalogue service. All existing SQL queries will need rewriting. 3 teams affected.",
     "Returns impact table: Technology HIGH, People MEDIUM, Data HIGH. Risk score: 7/10."),
]

for icon, name, when, why, how, demo, note in ba_tools:
    d.tool_entry(icon, name, when, why, how, demo, note, HEX_AMBER_BG, AMBER_LIGHT)

add_page_break(d.doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 9 — CHROME EXTENSION
# ─────────────────────────────────────────────────────────────────────────────
d.h1("9. Chrome Extension")
d.body(
    "The AiMitra Chrome Extension brings AI-powered test generation directly into the browser. It reads the live "
    "DOM of any webpage and generates complete, runnable test scripts using the actual CSS selectors, XPaths, "
    "and form structures found on the page — no generic placeholders."
)

d.h2("Features")
for f in [
    "DOM analysis: captures forms, inputs, buttons, links, ARIA roles, data-testid attributes",
    "4 frameworks: Playwright, Selenium, Cypress, WebdriverIO",
    "4 languages: Python, JavaScript, TypeScript, Java",
    "3 patterns: POM (Page Object Model), BDD/Gherkin, Simple function-based",
    "13 AI providers configurable — same providers as the main app",
    "Auto-injects content script into pages opened before extension was loaded",
]:
    d.bullet(f)

d.h2("Installation")
for step in [
    "Open Chrome and navigate to chrome://extensions/",
    "Enable 'Developer Mode' (toggle, top right)",
    "Click 'Load Unpacked'",
    "Select the chrome-extension/ folder from the AiMitra project",
    "The AiMitra icon appears in the Chrome toolbar",
]:
    d.bullet(step)

d.h2("Demo")
d.code_box("Navigate to any site", "https://automationpractice.pl/index.php?controller=authentication")
d.code_box("Configure and generate", "Framework: Playwright | Pattern: POM | Language: Python -> Click Generate")
d.body("Expected: A complete Python POM class with real selectors plus a test file with login happy path and negative test cases.")

d.divider()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 10 — PROVIDERS
# ─────────────────────────────────────────────────────────────────────────────
d.h1("10. AI Providers & Models")

providers = [
    ("Google Gemini", "Free", "gemini-2.0-flash, gemini-1.5-flash, gemma-3-27b", "Best free option for vision and large context"),
    ("Groq", "Free", "llama-3.3-70b, llama-3.1-8b, mixtral-8x7b", "Fastest inference — ideal for quick answers"),
    ("Cerebras", "Free", "llama-3.3-70b, qwen-3-32b", "Ultra-fast alternative to Groq"),
    ("OpenRouter", "Free*", "llama-3.3-70b-free, many others", "Aggregator — 200+ models, free tier available"),
    ("OpenAI", "Paid", "gpt-4o, gpt-4o-mini, o3-mini, o1", "Best overall quality, strong at code"),
    ("Anthropic", "Paid", "claude-opus-4-7, claude-sonnet-4-6, claude-haiku-4-5", "Best for long documents, precise reasoning"),
    ("Mistral AI", "Paid", "mistral-large, codestral, mistral-small", "Strong European alternative, good at code"),
    ("DeepSeek", "Paid", "deepseek-v3, deepseek-r1 (thinking)", "Excellent reasoning, very cost-effective"),
    ("xAI (Grok)", "Paid", "grok-3, grok-3-mini, grok-2", "Real-time web access in some modes"),
    ("Together AI", "Paid", "llama-3.3-70b, qwen, deepseek-r1", "Wide model selection, good batch pricing"),
    ("Perplexity", "Paid", "sonar-pro, sonar, r1-1776", "Best for web-search-augmented answers"),
    ("Fireworks AI", "Paid", "llama-3.3-70b, deepseek-r1, qwen", "Low latency, good for production use"),
    ("Cohere", "Paid", "command-r+, command-r", "Best for RAG and enterprise search workflows"),
]

table = d.doc.add_table(rows=1 + len(providers), cols=4)
table.style = "Table Grid"

# Header row
header_texts = ["Provider", "Tier", "Key Models", "Best For"]
header_widths = [Inches(1.3), Inches(0.6), Inches(2.3), Inches(2.1)]
for i, (txt, w) in enumerate(zip(header_texts, header_widths)):
    cell = table.cell(0, i)
    cell.width = w
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(txt)
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = WHITE
    shade_cell(cell, "3730A3")

# Data rows
for row_i, (pname, tier, models, best) in enumerate(providers):
    row = table.rows[row_i + 1]
    bg = "F5F3FF" if row_i % 2 == 0 else "FFFFFF"
    data = [pname, tier, models, best]
    for col_i, txt in enumerate(data):
        cell = row.cells[col_i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(8.5)
        if col_i == 1:
            run.font.color.rgb = EMERALD if "Free" in tier else AMBER
            run.font.bold = True
        else:
            run.font.color.rgb = DARK_TEXT
        shade_cell(cell, bg)

d.doc.add_paragraph()
d.divider()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 11 — ARCHITECTURE
# ─────────────────────────────────────────────────────────────────────────────
d.h1("11. Architecture & Data Flow")
d.body(
    "AiMitra uses a clean client-server architecture where the React frontend handles all UI state and the "
    "FastAPI backend acts as a secure proxy between the user and AI provider APIs."
)

d.h2("Request Flow — Non-Streaming")
steps_arch = [
    "User types a message and clicks Send",
    "React sends POST /api/chat with: message, mode, api_key, model, provider, temperature, history",
    "FastAPI validates the request (Pydantic ChatRequest model)",
    "prompt_router.py maps the mode to a (system_prompt, user_message) pair",
    "_pick_service() selects the correct backend service based on provider/model",
    "The service calls the AI provider's API (httpx async call)",
    "Response text is returned to FastAPI -> JSON response to React -> rendered as markdown",
]
for i, s in enumerate(steps_arch, 1):
    d.numbered(s, i)

d.h2("Request Flow — Streaming (SSE)")
d.body(
    "When streaming is enabled (default), the same request goes to POST /api/chat/stream. FastAPI returns a "
    "StreamingResponse wrapping an async generator. Each AI token is sent as 'data: {\"chunk\": \"...\"}' via SSE. "
    "React's native Fetch API reads the ReadableStream line-by-line and appends each chunk to the message in real time."
)

d.h2("Security Design")
for item in [
    "API keys are stored in browser localStorage — never on the server or in any database",
    "Keys are forwarded per-request in memory only — not logged, not cached server-side",
    "API key patterns are scrubbed from all error messages before returning to the frontend",
    "No user accounts required — zero PII collected",
]:
    d.bullet(item)

d.h2("Deployment Options")
for item in [
    "Single-service: FastAPI serves the built React app from backend/static/ (npm run build -> copy to static/)",
    "Free hosting: Render.com, Railway.app (backend + static), Vercel (frontend only)",
    "Docker: docker build -t aimitra . && docker run -p 8000:8000 aimitra",
]:
    d.bullet(item)

d.divider()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 12 — DEVELOPER
# ─────────────────────────────────────────────────────────────────────────────
d.h1("12. About the Developer", AMBER_LIGHT)

# Developer card
table_dev = d.doc.add_table(rows=1, cols=1)
table_dev.style = "Table Grid"
cell_dev = table_dev.cell(0, 0)
shade_cell(cell_dev, "F5F3FF")

p_name = cell_dev.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_name.paragraph_format.space_before = Pt(10)
run_name = p_name.add_run("Shiv Kant Kumar")
run_name.font.name  = "Calibri"
run_name.font.size  = Pt(20)
run_name.font.bold  = True
run_name.font.color.rgb = PURPLE_LIGHT

p_role = cell_dev.add_paragraph()
p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_role = p_role.add_run("Software Engineer  |  QA Enthusiast  |  AI Tools Builder")
run_role.font.name  = "Calibri"
run_role.font.size  = Pt(11)
run_role.font.color.rgb = AMBER_LIGHT

p_li_label = cell_dev.add_paragraph()
p_li_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_li_label.paragraph_format.space_before = Pt(6)
run_lil = p_li_label.add_run("Connect on LinkedIn:")
run_lil.font.size = Pt(10)
run_lil.font.color.rgb = DARK_TEXT

p_li_link = cell_dev.add_paragraph()
p_li_link.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_li2 = p_li_link.add_run("https://www.linkedin.com/in/shivkantkumar/")
run_li2.font.name  = "Calibri"
run_li2.font.size  = Pt(11)
run_li2.font.bold  = True
run_li2.font.color.rgb = BLUE

p_bio = cell_dev.add_paragraph()
p_bio.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_bio.paragraph_format.space_before = Pt(8)
p_bio.paragraph_format.space_after  = Pt(10)
run_bio = p_bio.add_run(
    "AiMitra was built to solve a real problem: QA engineers, BAs, and developers need\n"
    "domain-expert AI tools, not generic chatbots.\n"
    "Every tool was designed from daily work experience."
)
run_bio.font.name   = "Calibri"
run_bio.font.size   = Pt(9.5)
run_bio.font.italic = True
run_bio.font.color.rgb = GRAY_TEXT

d.doc.add_paragraph()
p_oss = d.doc.add_paragraph()
p_oss.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_oss = p_oss.add_run("AiMitra is open-source and free to use. Star it, fork it, build on it.")
run_oss.font.size = Pt(10)
run_oss.font.bold = True
run_oss.font.color.rgb = PURPLE_LIGHT

p_footer = d.doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_footer = p_footer.add_run("Document generated automatically from source.  Version 1.0  |  2025")
run_footer.font.size = Pt(8)
run_footer.font.italic = True
run_footer.font.color.rgb = GRAY_TEXT

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), "AiMitra_Documentation.docx")
d.save(out_path)
